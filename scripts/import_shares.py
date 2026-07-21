from __future__ import annotations

import argparse
import csv
import re
import shutil
import unicodedata
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "data" / "raw" / "分享"
ORGANIZED_DIR = ROOT / "data" / "processed" / "整理后的分享文章"
POSTS_DIR = ROOT / "src" / "content" / "posts"
REPORT_DIR = ROOT / "docs" / "内容整理报告"

COPYRIGHT_REVIEW = (
    "C.S.路易斯",
    "魔鬼家书",
    "以諾一書",
    "以诺一书",
    "Jesus loves Barabbas",
    "灵程指引",
)

THEOLOGY_REVIEW = (
    "88雅各",
    "91神的儿子，神的儿子们，人子",
)

THEOLOGY_APPROVED = (
    "88雅各",
    "91神的儿子，神的儿子们，人子",
)

OVERRIDES = {
    "圣经中麦子和稗子的意义是什么？": (
        "马太福音 13:24-30",
        "麦子和稗子的比喻：等候神最终的审判",
        "灵命成长",
    ),
    "100不贫穷也不富足": ("箴言 30:1-9", "不贫穷也不富足：亚古珥关于知足的祷告", "教会讲道"),
    "114主祷文的奥秘": ("马太福音 6:9-13", "主祷文：从天父关系理解祷告", "灵命成长"),
    "20220402 婚姻和教会": ("创世记 2:18-25", "婚姻与教会：在盟约中学习彼此相爱", "灵命成长"),
    "20220423_悔改": ("路加福音 24:47", "悔改：蒙赦免并回转归向神", "灵命成长"),
    "20220520_新人而非好人": ("约翰福音 3:1-15", "成为新人，而不只是成为好人", "灵命成长"),
    "20220628‗受苦的意义": ("约伯记 42:1-6", "受苦的意义：从约伯的经历认识神", "教会讲道"),
    "20220701_重新得力的秘诀": ("马太福音 11:28-30", "重新得力的秘诀：到基督里得安息", "灵命成长"),
    "20230503_假信仰与假平安": ("马太福音 7:15-23", "假信仰与假平安：从生命果子分辨真假", "教会讲道"),
    "33为什么耶稣告诉人们“去吧，从此不要再犯罪了”？": ("约翰福音 8:1-11", "去吧，从此不要再犯罪：恩典与圣洁的生命", "教会讲道"),
    "43三段被误解的圣经经文，以及神真正的应许": ("罗马书 8:28等", "三段常被误解的经文：在语境中理解神的应许", "教会讲道"),
    "47建在永恒根基上的生命——从两则比喻看属灵智慧": ("路加福音 16:19-31；马太福音 7:24-27", "建在永恒根基上的生命：从两则比喻学习属灵智慧", "教会讲道"),
    "77大卫和耶稣的祷告": ("诗篇 23:1；马太福音 6:9-13", "大卫与耶稣的祷告：学习信靠、敬拜与亲近神", "灵命成长"),
    "82保罗的刺": ("哥林多后书 12:7-10", "保罗身上的刺：在软弱中经历够用的恩典", "教会讲道"),
    "84聪明的童女": ("马太福音 25:1-13", "十个童女的比喻：警醒预备等候主", "教会讲道"),
    "85耶稣为何在地上写字": ("约翰福音 8:1-11", "耶稣为何在地上写字：恩典、审判与悔改", "教会讲道"),
    "86人生只有两条路": ("马太福音 7:13-14", "两条道路：选择通往生命的窄路", "教会讲道"),
    "87希伯来人，以色列人和犹太人": ("创世记 14:13；32:28", "希伯来人、以色列人与犹太人：名称与历史脉络", "教会讲道"),
    "88耶稣为何说她是狗": ("马太福音 15:21-28", "迦南妇人的信心：理解儿女与小狗的比喻", "教会讲道"),
    "88雅各": ("创世记 32:22-32", "雅各与神摔跤：从抓取到降服的生命转变", "教会讲道"),
    "91神的儿子，神的儿子们，人子": ("创世记 6:1-4；约伯记 1:6；但以理书 7:13-14", "神的众子、神的儿子与人子：三个称号的含义", "教会讲道"),
    "92大儿子的比喻": ("路加福音 15:25-32", "大儿子的比喻：隐藏在顺服外表下的疏离", "教会讲道"),
    "99雅各交叉双手": ("创世记 48:13-20", "雅各交叉双手祝福：神的主权超越人的次序", "教会讲道"),
    "为主而活 罗十四７，８": ("罗马书 14:7-8", "为主而活：无论生死都是属主的人", "灵命成长"),
    "从重生而成长": ("彼得前书 1:18-2:3", "从重生走向成长：渴慕纯净的灵奶", "灵命成长"),
    "以善胜恶顺服神": ("罗马书 12:19-21", "以善胜恶：把伸冤的主权交给神", "灵命成长"),
    "你是否真正得救": ("哥林多后书 13:5", "你是否真正得救：省察信心与生命果子", "灵命成长"),
    "信心": ("雅各书 2:14-26", "信心与行为：活的信心必结出果子", "灵命成长"),
    "和睦相处荣耀神": ("罗马书 12:14-18", "与人和睦：以祝福和善意回应冲突", "灵命成长"),
    "基督门徒的条件": ("路加福音 14:25-33", "作基督门徒的条件：计算跟随主的代价", "灵命成长"),
    "如何明辨神的声音": ("提摩太后书 3:16-17", "如何明辨神的声音：以圣经作为可靠准则", "灵命成长"),
    "百般试炼更坚信": ("雅各书 1:1-4", "在百般试炼中坚定信心", "灵命成长"),
    "神充分的奖赏": ("但以理书 3:1-30", "烈火窑中的信心：神同在就是充分的奖赏", "教会讲道"),
    "神完全的能力": ("哥林多后书 12:7-10", "神的能力在人的软弱上显得完全", "灵命成长"),
    "神阻挡骄傲的人": ("罗马书 11:16-24", "不可向旧枝夸口：在神的恩慈与严厉中警醒", "教会讲道"),
    "诗 32 悔改的操练": ("诗篇 32:1-11", "悔改的操练：在认罪与赦免中经历神", "灵命成长"),
    "重生的真义": ("约翰福音 3:3-8", "重生的真义：从圣灵而生的新生命", "灵命成长"),
    "承载恩典": ("罗马书 1:11-12", "承载恩典：在共同信心中彼此坚固", "灵命成长"),
    "寻求神的旨意": ("彼得后书 3:9；提摩太前书 2:3-4", "寻求神的旨意：从得救开始明白神的带领", "灵命成长"),
    "恩典": ("以弗所书 2:8-9", "恩典：白白得来的救恩", "灵命成长"),
    "列王纪上17_以利亚与撒勒法寡妇_信靠神的供应": (
        "列王纪上 17:1-24",
        "以利亚与撒勒法寡妇：信靠神的供应",
        "灵命成长",
    ),
    "创世记4_该隐与挪得之地_流离与救赎": (
        "创世记 4:1-24",
        "该隐与挪得之地：流离、罪的延续与救赎的盼望",
        "灵命成长",
    ),
    "20260721_使徒行传17章1-9节_帖撒罗尼迦在逼迫中持守福音的教会_JoelRyan_中文": (
        "使徒行传 17:1-9",
        "帖撒罗尼迦：在逼迫中持守福音的教会",
        "灵命成长",
    ),
    "20260721_帖撒罗尼迦前书5章17节_不住地祷告是什么意思_Ronnie_中文": (
        "帖撒罗尼迦前书 5:17",
        "不住地祷告是什么意思",
        "灵命成长",
    ),
    "20260721_帖撒罗尼迦前书4章16节_在基督里死了的人必先复活_Ronnie_中文": (
        "帖撒罗尼迦前书 4:16",
        "在基督里死了的人必先复活是什么意思",
        "灵命成长",
    ),
}

DESCRIPTION_OVERRIDES = {
    "圣经中麦子和稗子的意义是什么？":
        "耶稣借麦子和稗子的比喻说明善恶暂时并存、最终审判属于神，并呼召信徒忠心成长、结出生命的果子。",
    "88雅各":
        "本文根据创世记 32:22-32，思想雅各在雅博渡口与神相遇的经历，说明神如何破碎人靠自己抓取的生命，并带领人进入降服与祝福。",
    "91神的儿子，神的儿子们，人子":
        "本文结合创世记、约伯记与但以理书，梳理“神的众子”“神的儿子”与“人子”三个称号的不同含义，并帮助读者更清楚认识基督的身份。",
}

TEXT_REPLACEMENTS = {
    "创世纪": "创世记",
    "约伯纪": "约伯记",
    "钉S字J": "钉十字架",
    "人囗": "人口",
    "无庸置疑": "毋庸置疑",
    "想像": "想象",
    "彷佛": "仿佛",
    "于于神": "于神",
    "神。 神。": "神。",
    "整整一整夜": "整整一夜",
    "散住十二个支派": "散居各地的十二个支派",
    "这是圣经中唯一一次神在与人的对抗中输了。不是象征,不是比喻,是字面意义上的整夜摔跤,神自己承认,你得了胜。":
        "经文说雅各“与神与人较力,都得了胜”。这不是说神的能力不及雅各,而是神主动借着这场相遇破碎雅各原有的依靠,同时赐给他祝福和新的名字。",
    "这是圣经中唯一一次神在与人的对抗中输了。不是象征，不是比喻，是字面意义上的整夜摔跤，神自己承认，你得了胜。":
        "经文说雅各“与神与人较力，都得了胜”。这不是说神的能力不及雅各，而是神主动借着这场相遇破碎雅各原有的依靠，同时赐给他祝福和新的名字。",
    "所以,这里的神的儿子们,百分之百不可能是人类,只能是天使,是那些在人类被造之前就已经存在的灵体。这一点基本上没有什么争议。":
        "从创造大地的语境来看,许多解经者认为这里的“神的众子”是天使或属灵受造者。不过,相关称谓在不同经文中的解释仍需结合上下文谨慎判断。",
    "所以，这里的神的儿子们，百分之百不可能是人类，只能是天使，是那些在人类被造之前就已经存在的灵体。这一点基本上没有什么争议。":
        "从创造大地的语境来看，许多解经者认为这里的“神的众子”是天使或属灵受造者。不过，相关称谓在不同经文中的解释仍需结合上下文谨慎判断。",
    "在那个听话的长子身上藏着被99%的人都忽略的最致命的罪。":
        "在那个听话的长子身上,藏着许多读者容易忽略的属灵问题。",
    "使我们能百分之百顺服": "使我们学习全然顺服",
    "所以是百分之百神的工作,你所能做的只是百分之零":
        "所以救恩完全是神恩典的工作,人不能靠自己的功德赚取",
}

RISK_PHRASES = (
    "百分之百",
    "没有任何争议",
    "唯一一次",
    "神输了",
    "神自己承认",
    "绝对不可能",
    "一定就是",
    "99%",
)

BOOK_ALIASES = {
    "创": "创世记", "创世记": "创世记",
    "出": "出埃及记", "出埃及记": "出埃及记",
    "申": "申命记", "申命记": "申命记",
    "书": "约书亚记", "约书亚记": "约书亚记",
    "伯": "约伯记", "约伯记": "约伯记",
    "诗": "诗篇", "诗篇": "诗篇",
    "箴": "箴言", "箴言": "箴言",
    "太": "马太福音", "马太福音": "马太福音",
    "可": "马可福音", "马可福音": "马可福音",
    "路": "路加福音", "路加福音": "路加福音",
    "约": "约翰福音", "约翰福音": "约翰福音",
    "徒": "使徒行传", "使徒行传": "使徒行传",
    "罗": "罗马书", "罗马书": "罗马书",
    "林前": "哥林多前书", "哥林多前书": "哥林多前书",
    "林后": "哥林多后书", "哥林多后书": "哥林多后书",
    "加": "加拉太书", "加拉太书": "加拉太书",
    "弗": "以弗所书", "以弗所书": "以弗所书",
    "腓": "腓立比书", "腓立比书": "腓立比书",
    "西": "歌罗西书", "歌罗西书": "歌罗西书",
    "帖前": "帖撒罗尼迦前书", "帖撒罗尼迦前书": "帖撒罗尼迦前书",
    "帖后": "帖撒罗尼迦后书", "帖撒罗尼迦后书": "帖撒罗尼迦后书",
    "提前": "提摩太前书", "提摩太前书": "提摩太前书",
    "提后": "提摩太后书", "提摩太后书": "提摩太后书",
    "来": "希伯来书", "希伯来书": "希伯来书",
    "雅": "雅各书", "雅各书": "雅各书",
    "彼前": "彼得前书", "彼得前书": "彼得前书",
    "彼后": "彼得后书", "彼得后书": "彼得后书",
    "约一": "约翰一书", "约翰一书": "约翰一书",
    "启": "启示录", "启示录": "启示录",
}

def read_docx(path: Path) -> str:
    document = Document(path)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n\n".join(paragraphs)


def read_source(path: Path) -> str:
    if path.suffix.lower() == ".txt":
        return path.read_text(encoding="utf-8-sig")
    return read_docx(path)


def source_files() -> list[Path]:
    return sorted(
        path for path in SOURCE_DIR.iterdir()
        if path.is_file() and path.suffix.lower() in {".docx", ".txt"}
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Import one share source safely.")
    parser.add_argument("--source-file", help="Single share source file to import.")
    parser.add_argument("--dry-run", action="store_true", help="Preview without writing files.")
    parser.add_argument("--description", help="Manual frontmatter summary. Do not pass body excerpts or templates.")
    parser.add_argument(
        "--tags",
        help="Comma-separated SEO topic tags. Use 2-6 precise tags; the scripture book is added automatically when available.",
    )
    return parser.parse_args()


def resolve_source_file(value: str) -> Path:
    candidate = Path(value).expanduser()
    if not candidate.is_absolute():
        direct = (ROOT / candidate).resolve()
        if direct.exists():
            candidate = direct
        else:
            candidate = (SOURCE_DIR / value).resolve()
    else:
        candidate = candidate.resolve()
    try:
        candidate.relative_to(SOURCE_DIR.resolve())
    except ValueError:
        raise SystemExit(f"Share source file must be under {SOURCE_DIR}: {candidate}")
    if not candidate.is_file() or candidate.suffix.lower() not in {".docx", ".txt"}:
        raise SystemExit(f"Share source file not found or unsupported: {candidate}")
    return candidate


def load_csv_rows(path: Path, columns: int) -> dict[tuple[str, ...], tuple[str, ...]]:
    if not path.exists():
        return {}
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.reader(handle)
        next(reader, None)
        rows = (tuple(row) for row in reader if len(row) == columns)
        return {(row[0],): row for row in rows}


def normalize_text(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace("‗", " ").replace("﹕", ":").replace("﹣", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"\s+([，。！？；：、])", r"\1", text)
    text = re.sub(r"([，。！？；：、])(?=[^\s，。！？；：、”’）】])", r"\1", text)
    for old, new in TEXT_REPLACEMENTS.items():
        text = text.replace(old, new)
    return text.strip()


def source_date(path: Path) -> str:
    # Website display date must be the publication date, not the source file date,
    # filename date, or filesystem modified time.
    return datetime.now().strftime("%Y-%m-%d")


def clean_summary_title(path: Path) -> str:
    title = re.sub(r"^(?:20\d{6}[_ ]*|\d{1,3})", "", path.stem)
    title = re.sub(r"\s*副本$", "", title)
    title = unicodedata.normalize("NFKC", title)
    title = re.sub(
        r"\s*(?:经文|圣经)?\s*[:：]?\s*(?:罗马书|罗|约翰福音|约|诗篇|诗)?\s*\d+\s*[:：,，]\s*\d+(?:\s*[-—–－]\s*\d+)?\s*$",
        "",
        title,
    )
    title = re.sub(r"\s+", " ", title)
    return title.strip(" _-—:：,，。?？")


def parse_scripture(text: str) -> str:
    aliases = sorted(BOOK_ALIASES, key=len, reverse=True)
    book_pattern = "|".join(re.escape(alias) for alias in aliases)
    patterns = [
        rf"(?P<book>{book_pattern})\s*(?P<chapter>\d+)\s*[:：,，]\s*(?P<verses>\d+(?:\s*[-–—－]\s*\d+)?(?:\s*[;；]\s*\d+(?:\s*[-–—－]\s*\d+)?)*)",
        rf"(?P<book>{book_pattern})\s*第?\s*(?P<chapter>\d+)\s*章\s*(?P<verses>\d+(?:\s*[-–—－]\s*\d+)?)?\s*节?",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if not match:
            continue
        book = BOOK_ALIASES.get(match.group("book"), match.group("book"))
        chapter = match.group("chapter")
        verses = (match.groupdict().get("verses") or "").replace(" ", "")
        verses = verses.replace("–", "-").replace("—", "-").replace("－", "-").replace("；", ";")
        return f"{book} {chapter}:{verses}" if verses else f"{book} {chapter}"
    return ""


def find_scripture(title: str, body: str) -> str:
    from_title = parse_scripture(title)
    if from_title:
        return from_title

    blocks = [block.strip() for block in re.split(r"\n\s*\n", body) if block.strip()]
    labeled = [
        block for block in blocks[:30]
        if any(label in block[:30] for label in ("经文", "圣经", "主题经文", "读经"))
    ]
    for block in labeled:
        found = parse_scripture(block)
        if found:
            return found

    for block in blocks[:12]:
        found = parse_scripture(block)
        if found:
            return found
    return ""


def slugify(value: str) -> str:
    value = value.lower()
    value = re.sub(r"[\s_：:;；,，.。()（）\[\]【】/\\|｜?？\"“”]+", "-", value)
    return re.sub(r"-+", "-", value).strip("-") or "share"


def yaml_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace('"', '\\"')


def strip_leading_title(text: str, source: Path) -> str:
    first, separator, remainder = text.partition("\n\n")
    if separator and first.strip(" ?？") == source.stem.strip(" ?？"):
        return remainder.lstrip()
    return text


def validate_manual_description(description: str, body: str) -> None:
    normalized = re.sub(r"\s+", " ", description).strip()
    placeholders = {"NEEDS_METADATA：请人工阅读文章后补充大意摘要。", "NEEDS_DESCRIPTION_REVIEW", "TODO", "TBD", "待补", "暂无"}
    if not normalized:
        raise SystemExit("Missing manual description. Add --description or DESCRIPTION_OVERRIDES entry.")
    if normalized in placeholders:
        raise SystemExit("Description is a placeholder; add a manual summary before publishing.")
    if len(normalized) < 30:
        raise SystemExit("Description is too short; add a complete manual summary.")
    if not re.search(r"[。！？.!?]$", normalized):
        raise SystemExit("Description must be a complete sentence ending with punctuation.")
    body_head = re.sub(r"\s+", " ", body[:300]).strip()
    if body_head and normalized[:40] in body_head:
        raise SystemExit("Description appears copied from the body opening; write a manual summary instead.")


def build_seo_tags(raw_tags: str | None, scripture: str) -> list[str]:
    if not raw_tags:
        raise SystemExit(
            "Missing SEO tags. Add --tags with 2-6 precise topic tags; "
            "include core people, places, doctrines, or applications rather than category labels."
        )
    supplied = [tag.strip() for tag in re.split(r"[,，]", raw_tags) if tag.strip()]
    generic = {"分享", "灵命成长", "文章", "信仰", "基督教"}
    if any(tag in generic for tag in supplied):
        raise SystemExit("SEO tags must be topical; remove generic category/process labels such as 分享 or 灵命成长.")
    book = scripture.split(" ", 1)[0].strip() if scripture else ""
    tags: list[str] = []
    for tag in ([book] if book else []) + supplied:
        if tag and tag not in tags:
            tags.append(tag)
    if not 2 <= len(tags) <= 6:
        raise SystemExit(f"SEO tags must total 2-6 after scripture-book normalization; got {len(tags)}: {tags}")
    return tags


def sentence_aware_description(body: str, fallback: str) -> str:
    blocks = [block.strip() for block in re.split(r"\n\s*\n", body) if block.strip()]
    sentences: list[str] = []
    for block in blocks[:8]:
        text = re.sub(r"\s+", " ", block).strip()
        if not text or text in {"⸻", "---"}:
            continue
        if len(text) <= 30 and not re.search(r"[。！？.!?]$", text):
            continue
        parts = re.findall(r".+?[。！？.!?](?=\s|$)|.+$", text)
        for part in parts:
            sentence = part.strip()
            if len(sentence) >= 24:
                sentences.append(sentence)
            if len("".join(sentences)) >= 80:
                return "".join(sentences)
    return "".join(sentences) or fallback


def format_body(text: str) -> str:
    blocks = []
    for block in re.split(r"\n\s*\n", text):
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if not lines:
            continue
        joined = " ".join(lines)
        if joined in {"⸻", "---"}:
            continue
        joined = re.sub(r"^#{1,6}\s+", "", joined).strip()
        if joined.endswith(("?", "？")) and len(joined) <= 40:
            blocks.append(f"## {joined}")
        elif re.match(r"^(?:[一二三四五六七八九十]+、|\d+[.、])", joined) and len(joined) <= 80:
            blocks.append(f"## {joined}")
        elif joined.startswith(("经文:", "经文：", "读经:", "读经：", "圣经:", "圣经：")):
            blocks.append(f"> {joined}")
        else:
            blocks.append(joined)
    return "\n\n".join(blocks)


def main() -> None:
    args = parse_args()
    if not args.source_file:
        raise SystemExit(
            "import_shares.py no longer imports all share files by default. "
            "Use --source-file to select one file; add --dry-run to preview."
        )
    target_source = resolve_source_file(args.source_file)
    if not args.dry_run:
        ORGANIZED_DIR.mkdir(exist_ok=True)
        POSTS_DIR.mkdir(parents=True, exist_ok=True)
        REPORT_DIR.mkdir(exist_ok=True)
    sources = [target_source]
    catalog_path = REPORT_DIR / "分享文章目录.csv"
    copyright_path = REPORT_DIR / "待确认版权.csv"
    missing_path = REPORT_DIR / "待补经文.csv"
    risk_path = REPORT_DIR / "待神学事实复核.csv"
    catalog_rows = load_csv_rows(catalog_path, 6)
    copyright_rows = load_csv_rows(copyright_path, 2)
    missing_rows = load_csv_rows(missing_path, 4)
    risk_rows = load_csv_rows(risk_path, 3)
    imported = 0
    for source in sources:
        if any(keyword.lower() in source.name.lower() for keyword in COPYRIGHT_REVIEW):
            copyright_rows[(source.name,)] = (source.name, "来源或版权需要确认，暂不发布")
            continue

        body = strip_leading_title(normalize_text(read_source(source)), source)
        summary = clean_summary_title(source)
        override = OVERRIDES.get(source.stem)
        if override:
            scripture, summary, _previous_category = override
        else:
            scripture = find_scripture(summary, body)
        category = "灵命成长"
        date = source_date(source)
        published_at = datetime.now(ZoneInfo("Asia/Tokyo")).isoformat(timespec="seconds")
        title = f"{scripture}｜{summary}" if scripture else summary
        reviewed = False
        description = args.description or DESCRIPTION_OVERRIDES.get(
            source.stem,
            "NEEDS_METADATA：请人工阅读文章后补充大意摘要。",
        )
        validate_manual_description(description, body)
        tags = build_seo_tags(args.tags, scripture)
        slug = f"{date}-{slugify(title)}"

        markdown = (
            "---\n"
            f'title: "{yaml_escape(title)}"\n'
            f'description: "{yaml_escape(description)}"\n'
            f"date: {date}\n"
            f"publishedAt: {published_at}\n"
            "tags: [" + ", ".join(f'"{yaml_escape(tag)}"' for tag in tags) + "]\n"
            f'category: "{yaml_escape(category)}"\n'
            f'scripture: "{yaml_escape(scripture)}"\n'
            'author: "Ronnie"\n'
            f"reviewed: {str(reviewed).lower()}\n"
            f'source: "data/raw/分享/{yaml_escape(source.name)}"\n'
            "---\n\n"
            + format_body(body)
            + "\n"
        )
        output = ORGANIZED_DIR / f"{slug}.md"
        post_output = POSTS_DIR / output.name
        existing_targets = [path for path in (output, post_output) if path.exists()]
        if existing_targets:
            details = "\n".join(f"- {path}" for path in existing_targets)
            raise SystemExit(
                "Target file already exists; stopping to avoid overwriting existing share posts.\n"
                f"{details}"
            )
        print(f"Source file: {source.name}")
        print(f"Title: {title}")
        print(f"Slug: {slug}")
        print(f"Processed target: {output}")
        print(f"Post target: {post_output}")
        if args.dry_run:
            print("Dry-run: no files written.")
        else:
            output.write_text(markdown, encoding="utf-8")
        requires_theology_review = (
            any(keyword.lower() in source.name.lower() for keyword in THEOLOGY_REVIEW)
            and not any(keyword.lower() in source.name.lower() for keyword in THEOLOGY_APPROVED)
        )
        publishable = bool(scripture) and not requires_theology_review
        if publishable and not args.dry_run:
            shutil.copyfile(output, post_output)
        status = "待神学复核" if requires_theology_review else ("待补经文" if not scripture else "已完成基础校订")
        row = (source.name, date, scripture, category, title, status)
        catalog_rows[(source.name,)] = row
        if scripture:
            missing_rows.pop((source.name,), None)
        else:
            missing_rows[(source.name,)] = (source.name, date, category, title)
        risk_rows.pop((source.name,), None)
        for phrase in RISK_PHRASES:
            if phrase in body:
                snippet_start = max(body.find(phrase) - 80, 0)
                snippet = re.sub(r"\s+", " ", body[snippet_start:snippet_start + 240])
                risk_rows[(source.name,)] = (source.name, phrase, snippet)
                break
        imported += 1

    if not args.dry_run:
        with catalog_path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(("源文件", "日期", "经文", "分类", "发布标题", "校订状态"))
            writer.writerows(catalog_rows.values())

        with copyright_path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(("源文件", "原因"))
            writer.writerows(copyright_rows.values())

        with missing_path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(("源文件", "日期", "分类", "暂定标题"))
            writer.writerows(missing_rows.values())

    if not args.dry_run:
        audit_path = REPORT_DIR / "分享文章审计摘录.md"
        audit_text = audit_path.read_text(encoding="utf-8") if audit_path.exists() else "# 分享文章审计摘录\n"
        audit_sections = []
        for source in sources:
            if any(keyword.lower() in source.name.lower() for keyword in COPYRIGHT_REVIEW):
                continue
            if f"## {source.name}\n" in audit_text:
                continue
            body = normalize_text(read_source(source))
            blocks = [block.strip() for block in re.split(r"\n\s*\n", body) if block.strip()]
            audit_lines = [f"## {source.name}", ""]
            for block in blocks[:10]:
                audit_lines.append(block[:500])
                audit_lines.append("")
            audit_sections.append("\n".join(audit_lines))
        if audit_sections:
            audit_text = audit_text.rstrip() + "\n\n" + "\n\n".join(
                section.rstrip() for section in audit_sections
            ) + "\n"
            audit_path.write_text(audit_text, encoding="utf-8")

        with risk_path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.writer(handle)
            writer.writerow(("源文件", "风险表达", "上下文摘录"))
            writer.writerows(risk_rows.values())

    print(f"Imported shares: {imported}")
    print(f"Copyright review: {len(copyright_rows)}")
    print(f"Missing scripture: {len(missing_rows)}")


if __name__ == "__main__":
    main()
